from __future__ import annotations
import platform
import os
import re
import argparse
import requests
import faiss
import numpy as np
from PIL import Image
from sentence_transformers import SentenceTransformer
from typing import List, Dict, Any, Iterable, Tuple

import fitz
import docx

# nexaai imports
from nexaai.llm import LLM as NexaPyLLM, GenerationConfig
from nexaai.common import ModelConfig, ChatMessage
from nexaai.embedder import Embedder, EmbeddingConfig
from nexaai.rerank import Reranker, RerankConfig

from langchain_core.language_models.llms import LLM
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.schema.runnable import RunnableLambda


# Nexa config
DEFAULT_MODEL = "NexaAI/Qwen3-VL-4B-Instruct-GGUF"

if platform.system() == "Windows":
    DEFAULT_EMBED_MODEL = "djuna/jina-embeddings-v2-small-en-Q5_K_M-GGUF"
    EMBEDDING_PLUGIN_ID = "cpu_gpu"
elif platform.system() == "Darwin": 
    DEFAULT_EMBED_MODEL = "NexaAI/jina-v2-fp16-mlx"
    EMBEDDING_PLUGIN_ID = "mlx"
else:
    DEFAULT_EMBED_MODEL = "djuna/jina-embeddings-v2-small-en-Q5_K_M-GGUF"
    EMBEDDING_PLUGIN_ID = "cpu_gpu"
    
def call_nexa_chat(model: str, messages: list):
    """
    Non-streaming conversation with local LLM
    """
    m_cfg = ModelConfig()
    llm = NexaPyLLM.from_(model, plugin_id="mlx", m_cfg=m_cfg)
    prompt = llm.apply_chat_template(messages)
    
    return llm.generate(prompt, g_cfg=GenerationConfig(max_tokens=512))


def stream_nexa_chat_messages(model: str, messages: list):
    """
    streaming conversation with local LLM
    """
    m_cfg = ModelConfig()
    llm = NexaPyLLM.from_(model, plugin_id="cpu_gpu", m_cfg=m_cfg)
    prompt = llm.apply_chat_template(messages)
    
    for piece in llm.generate_stream(prompt, g_cfg=GenerationConfig(max_tokens=512)):
        yield piece

def call_nexa_embeddings(embed_model: str, inputs: List[str]) -> List[List[float]]:
    """
    Call nexa embeddding.
    
    Args:
        embed_model: Embedding model name
        inputs: List of text strings to embed
        model_folder: model folder path
        
    Returns:
        List[List[float]]: List of embedding vectors aligned to input order
    """
    if not inputs:
        return []
    
    out: List[List[float]] = []
    
    embedder = Embedder.from_(name_or_path=embed_model, plugin_id=EMBEDDING_PLUGIN_ID)

    # Process in batches to avoid large payloads
    BATCH_SIZE = 64
    for i in range(0, len(inputs), BATCH_SIZE):
        batch = inputs[i:i+BATCH_SIZE]
        batch_size = len(batch)
        embeddings = embedder.generate(
        texts=batch, config=EmbeddingConfig(batch_size=batch_size))
        
        for embedding in embeddings:
            out.append(embedding.tolist())
           
    return out


class NexaLLM(LLM):
    """A minimal LangChain LLM adapter that calls Nexa's OpenAI-style endpoints."""
    model: str = DEFAULT_MODEL

    def _call(self, prompt: str, **kwargs: Any) -> str:
        messages = [
            {"role": "user", "content": prompt},
        ]
        return call_nexa_chat(self.model, messages=messages)

    @property
    def _llm_type(self) -> str:
        return f"nexa:{self.model}"


# File loaders
def load_txt(path: str) -> str:
    """Read UTF-8 text; fall back to latin-1 if needed."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, "r", encoding=enc, errors="ignore") as f:
                return f.read()
        except Exception:
            continue
    return ""

def load_pdf(path: str) -> str:
    """Extract text from PDF using PyMuPDF (page by page)."""
    text_parts: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)

def load_docx(path: str) -> str:
    """Extract text from .docx using python-docx."""
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs]
    return "\n".join(paras)

def normalize_ws(s: str) -> str:
    """Collapse consecutive whitespace for cleaner chunks."""
    return re.sub(r"[ \t\u3000]+", " ", s).strip()

def yield_files(root: str, exts: Tuple[str, ...] = (".txt", ".pdf", ".docx")) -> Iterable[str]:
    """Iterate all files (recursively) under root with given extensions."""
    for base, _, files in os.walk(root):
        for name in files:
            if name.lower().endswith(exts):
                yield os.path.join(base, name)

def load_file(path: str) -> str:
    """Dispatch by extension to proper loader."""
    lower = path.lower()
    try:
        if lower.endswith(".txt"):
            return load_txt(path)
        if lower.endswith(".pdf"):
            return load_pdf(path)
        if lower.endswith(".docx"):
            return load_docx(path)
    except Exception as e:
        print(f"[warn] Failed to read {path}: {e}")
        return ""
    return ""

def yield_images(root: str, exts: Tuple[str, ...] = (".png", ".jpg", ".jpeg", ".webp", ".bmp")) -> List[str]:
    """Collect image file paths under root (recursive)."""
    paths: List[str] = []
    for base, _, files in os.walk(root):
        for name in files:
            if name.lower().endswith(exts):
                # Use forward slashes to avoid JSON escape headaches on Windows
                p = os.path.abspath(os.path.join(base, name)).replace("\\", "/")
                paths.append(p)
    return paths


def build_image_index(image_paths: List[str],
                      model_name: str = "clip-ViT-B-32") -> tuple[faiss.IndexFlatIP, List[str], SentenceTransformer]:
    """
    Build a FAISS index for images using a CLIP model (cross-modal).
    Returns: (faiss_index, paths, clip_model)
    """
    if not image_paths:
        return None, [], None  # type: ignore

    # Load CLIP model (works for both text and images)
    clip_model = SentenceTransformer(model_name)

    # Encode images -> L2-normalized embeddings so inner product == cosine
    embs = []
    for p in image_paths:
        try:
            img = Image.open(p).convert("RGB")
            em = clip_model.encode(img, convert_to_numpy=True, normalize_embeddings=True)
            embs.append(em)
        except Exception as e:
            print(f"[warn] failed to embed image {p}: {e}")
            embs.append(None)

    # Filter out failed ones
    kept_paths, kept_embs = [], []
    for p, e in zip(image_paths, embs):
        if e is not None:
            kept_paths.append(p)
            kept_embs.append(e)
    if not kept_embs:
        return None, [], clip_model  # type: ignore

    mat = np.vstack(kept_embs).astype("float32")
    index = faiss.IndexFlatIP(mat.shape[1])   # cosine via inner product on normalized vectors
    index.add(mat)
    print(f"[info] Image index built: {len(kept_paths)} images, dim={mat.shape[1]}")
    return index, kept_paths, clip_model

def retrieve_topk_images(query: str, k: int,
                        index: faiss.IndexFlatIP,
                        paths: List[str],
                        clip_model: SentenceTransformer) -> List[str]:
    """
    Text->image retrieval: encode text with CLIP and search the image FAISS.
    """
    if index is None or clip_model is None or not paths:
        return []
    q = clip_model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
    D, I = index.search(q, min(k, len(paths)))
    return [paths[i] for i in I[0] if 0 <= i < len(paths)]


# Chunking & retriever
def build_chunks_from_folder(folder: str,
                            chunk_size: int = 1000,
                            chunk_overlap: int = 150) -> List[Document]:
    """Load files from folder, split into chunks, attach metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", "。", "；", "，", " ", ""],
    )

    docs: List[Document] = []
    for path in yield_files(folder):
        raw = load_file(path)
        raw = normalize_ws(raw)
        if not raw:
            continue

        chunks = splitter.split_text(raw)
        for i, ch in enumerate(chunks):
            docs.append(
                Document(
                    page_content=ch,
                    metadata={
                        "source": os.path.abspath(path),
                        "chunk_index": i,
                        "total_chars": len(raw),
                    },
                )
            )
    return docs

class _ServerEmbeddingRetriever:
    """
    Minimal retriever that queries Nexa /v1/embeddings for both index-building
    and query embedding, and searches with FAISS (cosine via inner product).
    """
    def __init__(self, texts: List[str], metas: List[dict], k: int, embed_model: str):
        self.embed_model = embed_model
        self.k = k
        self.texts = texts
        self.metas = metas

        # Build matrix via server-side embeddings
        vecs = call_nexa_embeddings(self.embed_model, self.texts)
        mat = np.array(vecs, dtype=np.float32)
        # Normalize rows for cosine inner product
        norms = np.linalg.norm(mat, axis=1, keepdims=True) + 1e-8
        mat = mat / norms

        self.index = faiss.IndexFlatIP(mat.shape[1])
        self.index.add(mat)

    def get_relevant_documents(self, query: str) -> List[Document]:
        q_vec = call_nexa_embeddings(self.embed_model, [query])[0]
        q = np.array(q_vec, dtype=np.float32)
        q = q / (np.linalg.norm(q) + 1e-8)
        D, I = self.index.search(q[np.newaxis, :], self.k)
        docs: List[Document] = []
        for i in I[0]:
            if 0 <= i < len(self.texts):
                docs.append(Document(page_content=self.texts[i], metadata=self.metas[i]))
        return docs

def build_retriever(docs: List[Document], k: int = 5, embed_model: str = DEFAULT_EMBED_MODEL):
    """Create FAISS retriever using server-side embeddings; attach metadata."""
    texts = [d.page_content for d in docs]
    metas = [d.metadata for d in docs]
    if not texts:
        return None
    return _ServerEmbeddingRetriever(texts, metas, k=k, embed_model=embed_model)


# Prompt template
SYSTEM_TEMPLATE = """
You are a careful assistant. Use ONLY the provided context to answer.
If the answer is not contained in the context, say you don't know and ask for more info.

<context>
{context}
</context>
"""

prompt = ChatPromptTemplate.from_messages(
    [
        ("system", SYSTEM_TEMPLATE),
        MessagesPlaceholder("messages"),
    ]
)

def build_chain(retriever, model_name: str):
    """Compose retrieval -> stuff documents -> Nexa LLM."""
    llm = NexaLLM(model=model_name)
    doc_chain = create_stuff_documents_chain(llm, prompt)

    def add_context(params: Dict[str, Any]) -> Dict[str, Any]:
        q = params["messages"][-1].content
        docs = retriever.get_relevant_documents(q) if retriever else []
        return {"messages": params["messages"], "context": docs}

    chain = RunnableLambda(add_context) | doc_chain
    return chain


# CLI
def main():
    ap = argparse.ArgumentParser(description="Local-files RAG with Nexa python binding")
    ap.add_argument("--data", default="./docs", help="Folder containing txt/pdf/docx (recursive).")
    ap.add_argument("--k", type=int, default=5, help="Top-k documents to retrieve.")
    ap.add_argument("--chunk_size", type=int, default=1000, help="Chunk size for splitting.")
    ap.add_argument("--chunk_overlap", type=int, default=150, help="Chunk overlap for splitting.")
    ap.add_argument("--model", default=DEFAULT_MODEL, help="Nexa model name or alias.")
    ap.add_argument("--embed_model", default=DEFAULT_EMBED_MODEL, help="Embedding model served by Nexa /v1/embeddings")
    args = ap.parse_args()

    if not os.path.exists(args.data):
        os.makedirs(args.data)
        print(f"[info] Created empty data folder: {args.data}")
        
    print(f"[info] Loading files from: {args.data}")
    docs = build_chunks_from_folder(args.data, args.chunk_size, args.chunk_overlap)
    if not docs:
        print("[error] No documents loaded. Check your --data path and file types.")
        return
    print(f"[info] Built {len(docs)} chunks.")

    # Build image index once (from ./docs)
    img_paths_all = yield_images(args.data)
    img_index, img_paths_kept, clip_model = build_image_index(img_paths_all)

    retriever = build_retriever(docs, k=args.k, embed_model=args.embed_model)
    chain = build_chain(retriever, model_name=args.model)

    print(f"[info] Ready. Using model={args.model}")
    print("Type your question (or just press Enter to quit):")

    while True:
        try:
            q = input("[user] ").strip()
            if not q:
                break

            # Hot-reload index on demand
            if q.lower() == ":reload":
                print("[info] Rebuilding index ...")
                docs = build_chunks_from_folder(args.data, args.chunk_size, args.chunk_overlap)
                retriever = build_retriever(docs, k=args.k, embed_model=args.embed_model)
                print(f"[info] Rebuilt. Chunks: {len(docs)}")
                continue

            # Retrieval only (no LLM call here)
            ctx_docs = retriever.get_relevant_documents(q) if retriever else []

            # Show retrieved evidence
            print("\n[retrieved]")
            for i, d in enumerate(ctx_docs, start=1):
                src = d.metadata.get("source", "")
                idx = d.metadata.get("chunk_index", -1)
                snippet = d.page_content[:160].replace("\n", " ")
                print(f"  {i}. {os.path.basename(src)}#chunk{idx}: {snippet}...")

            # Streaming generation
            context_text = "\n\n".join([d.page_content for d in ctx_docs])
            IMG_TOPK = 1
            topk_imgs = retrieve_topk_images(q, IMG_TOPK, img_index, img_paths_kept, clip_model)

            # For logging
            if topk_imgs:
                print("\n[images-selected]")
                for i, p in enumerate(topk_imgs, 1):
                    print(f"  {i}. {p}")

            img_contents = [{"type": "image_url", "image_url": {"url": p}} for p in topk_imgs]

            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are a careful assistant. Use ONLY the provided context to answer.\n\n"
                        f"<context>\n{context_text}\n</context>"
                    ),
                },
                {
                    "role": "user",
                    "content": [{"type": "text", "text": q}] + img_contents
                },
            ]

            print("\n[assistant]", end="", flush=True)
            try:
                # stream_nexa_chat_messages yields incremental text pieces
                for piece in stream_nexa_chat_messages(args.model, messages):
                    print(piece, end="", flush=True)
                print()  # newline after stream ends
            except requests.HTTPError as e:
                # Fallback to non-stream single-shot completion (build a single prompt string)
                print(f"\n[warn] streaming failed, fallback to non-stream. Reason: {e}")
                prompt_text = (
                    "You are a careful assistant. Use ONLY the provided context to answer.\n\n"
                    f"<context>\n{context_text}\n</context>\n\n"
                    f"Question: {q}"
                )
                try:
                    messages = [{
                        "role": "user",
                        "content": prompt_text
                    }]
                    nonstream_answer = call_nexa_chat(args.model, messages=messages)
                    print("\n[assistant]")
                    print(nonstream_answer)
                except Exception as e2:
                    print(f"[error] Non-stream request also failed: {e2}")

        except KeyboardInterrupt:
            print("\n[info] Bye.")
            break

if __name__ == "__main__":
    main()
